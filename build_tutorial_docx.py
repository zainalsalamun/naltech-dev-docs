from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "tutorial-flutter-dasar-pemula.md"
OUTPUT = ROOT / "docs" / "tutorial-flutter-dasar-pemula.docx"


COLORS = {
    "blue": RGBColor(0x2E, 0x74, 0xB5),
    "dark_blue": RGBColor(0x1F, 0x4D, 0x78),
    "ink": RGBColor(0x1F, 0x29, 0x37),
    "muted": RGBColor(0x5B, 0x67, 0x72),
    "code_fill": "F4F6F9",
    "border": "D9E2EC",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9E2EC", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = COLORS["dark_blue"]
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = COLORS["muted"]
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, COLORS["blue"]),
        ("Heading 2", 13, 14, 7, COLORS["blue"]),
        ("Heading 3", 12, 10, 5, COLORS["dark_blue"]),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.text = "Tutorial Flutter Dasar untuk Pemula"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = COLORS["muted"]


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Tutorial Flutter Dasar untuk Pemula")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Panduan belajar lengkap untuk developer Flutter pemula: setup, Dart dasar, widget, layout, state, navigasi, form, API, penyimpanan sederhana, dan mini project."
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_width(table)
    for cell in table.rows[0].cells:
        set_cell_margins(cell)
        set_cell_border(cell, "D9E2EC")
    table.rows[0].cells[0].text = "Format"
    table.rows[0].cells[1].text = "Reference guide / modul belajar"
    set_cell_shading(table.rows[0].cells[0], "E8EEF5")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Cara pakai dokumen ini")
    r.bold = True
    r.font.color.rgb = COLORS["dark_blue"]
    for item in [
        "Baca satu bab pendek, lalu langsung jalankan contoh kodenya.",
        "Gunakan checklist akhir untuk mengukur progress belajar.",
        "Kerjakan mini project catatan belajar sebelum lanjut ke state management kompleks.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["code_fill"])
    set_cell_border(cell, COLORS["border"])
    set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.rstrip().splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x11, 0x27, 0x35)
    doc.add_paragraph()


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLORS["dark_blue"]
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(line.strip() for line in buffer).strip()
    if text:
        p = doc.add_paragraph()
        add_inline_markdown(p, text)
    buffer.clear()


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(" ")
    run.font.size = Pt(1)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EC")
    border.append(bottom)
    p_pr.append(border)


def convert_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    paragraph_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    skip_title = True

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buffer))
                code_buffer.clear()
                in_code = False
            else:
                flush_paragraph(doc, paragraph_buffer)
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            flush_paragraph(doc, paragraph_buffer)
            continue

        if line.strip() == "---":
            flush_paragraph(doc, paragraph_buffer)
            add_horizontal_rule(doc)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush_paragraph(doc, paragraph_buffer)
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if skip_title and level == 1:
                skip_title = False
                continue
            skip_title = False
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            doc.add_paragraph(text, style=style)
            continue

        bullet = re.match(r"^\s*-\s+(.+)$", line)
        checklist = re.match(r"^\s*-\s+\[( |x|X)\]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)

        if checklist:
            flush_paragraph(doc, paragraph_buffer)
            mark = "☑" if checklist.group(1).lower() == "x" else "☐"
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, f"{mark} {checklist.group(2)}")
            continue

        if bullet:
            flush_paragraph(doc, paragraph_buffer)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet.group(1))
            continue

        if numbered:
            flush_paragraph(doc, paragraph_buffer)
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, numbered.group(1))
            continue

        paragraph_buffer.append(line)

    flush_paragraph(doc, paragraph_buffer)
    if code_buffer:
        add_code_block(doc, "\n".join(code_buffer))


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    convert_markdown(doc, markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
